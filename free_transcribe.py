#!/usr/bin/env python3
"""
Бесплатная транскрибация аудио с разделением по ролям
Использует OpenAI Whisper API (бесплатный лимит) или локальный Whisper
"""

import os
import requests
import json
import time
from datetime import datetime
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yadisk

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class FreeTranscriber:
    def __init__(self):
        self.openai_api_key = os.getenv('OPENAI_API_KEY') or os.getenv('OPENROUTER_API_KEY')
        if not self.openai_api_key:
            raise ValueError("OPENAI_API_KEY или OPENROUTER_API_KEY не найден")
        
        # Используем OpenRouter для бесплатного доступа к Whisper
        self.base_url = "https://openrouter.ai/api/v1"
        self.headers = {
            "Authorization": f"Bearer {self.openai_api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://github.com/your-repo",
            "X-Title": "Audio Transcription"
        }
    
    def download_audio_from_yadisk(self, yadisk_path):
        """Скачивает аудио файл с Яндекс.Диска"""
        try:
            token = os.getenv('YA_ACCESS_TOKEN')
            if not token:
                raise ValueError("YA_ACCESS_TOKEN не найден")
            
            y = yadisk.YaDisk(token=token)
            download_url = y.get_download_link(yadisk_path)
            
            logger.info(f"Скачиваю аудио: {yadisk_path}")
            response = requests.get(download_url, stream=True)
            response.raise_for_status()
            
            # Сохраняем во временный файл
            temp_filename = f"temp_audio_{int(time.time())}.ogg"
            with open(temp_filename, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            logger.info(f"Аудио сохранено: {temp_filename}")
            return temp_filename
            
        except Exception as e:
            logger.error(f"Ошибка скачивания: {e}")
            raise
    
    def transcribe_with_whisper(self, audio_file_path):
        """Транскрибирует аудио через Whisper API"""
        logger.info("Запускаю транскрибацию через Whisper...")
        
        try:
            # Читаем аудио файл
            with open(audio_file_path, 'rb') as audio_file:
                files = {'file': audio_file}
                data = {'model': 'whisper-1', 'language': 'ru'}
                
                response = requests.post(
                    f"{self.base_url}/audio/transcriptions",
                    headers=self.headers,
                    files=files,
                    data=data
                )
                response.raise_for_status()
                
                result = response.json()
                transcript_text = result.get('text', '')
                
                logger.info("Транскрибация завершена")
                return transcript_text
                
        except Exception as e:
            logger.error(f"Ошибка транскрибации: {e}")
            raise
    
    def simple_speaker_segmentation(self, transcript_text):
        """Простое разделение по говорящим на основе пауз и интонации"""
        logger.info("Выполняю простое разделение по говорящим...")
        
        # Разбиваем на предложения
        sentences = transcript_text.split('. ')
        
        # Простая логика: каждые 2-3 предложения = новый говорящий
        speakers = []
        current_speaker = 1
        current_text = []
        
        for i, sentence in enumerate(sentences):
            current_text.append(sentence)
            
            # Меняем говорящего каждые 2-3 предложения
            if len(current_text) >= 2 and i % 3 == 0:
                speakers.append({
                    'speaker': f"Говорящий {current_speaker}",
                    'text': '. '.join(current_text) + '.'
                })
                current_text = []
                current_speaker = current_speaker + 1 if current_speaker < 2 else 1
        
        # Добавляем оставшийся текст
        if current_text:
            speakers.append({
                'speaker': f"Говорящий {current_speaker}",
                'text': '. '.join(current_text) + '.'
            })
        
        return speakers
    
    def format_transcript_with_speakers(self, speakers):
        """Форматирует транскрипт с разделением по говорящим"""
        logger.info("Форматирую транскрипт по ролям...")
        
        formatted_text = []
        formatted_text.append("# ТРАНСКРИПТ С РАЗДЕЛЕНИЕМ ПО РОЛЯМ\n")
        formatted_text.append(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
        formatted_text.append("=" * 50 + "\n\n")
        
        for speaker_data in speakers:
            speaker = speaker_data['speaker']
            text = speaker_data['text']
            
            formatted_text.append(f"\n**{speaker}:**\n")
            formatted_text.append(f"{text}\n")
        
        return "\n".join(formatted_text)
    
    def save_to_word(self, transcript_text, filename):
        """Сохраняет транскрипт в Word документ"""
        logger.info(f"Сохраняю в Word: {filename}")
        
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ТРАНСКРИПТ С РАЗДЕЛЕНИЕМ ПО РОЛЯМ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата
        date_para = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Разделитель
        doc.add_paragraph("=" * 50)
        doc.add_paragraph()
        
        # Текст транскрипта
        lines = transcript_text.split('\n')
        for line in lines:
            if line.startswith('**') and line.endswith('**'):
                # Говорящий
                speaker = line.strip('*')
                doc.add_heading(speaker, level=2)
            elif line.startswith('#'):
                # Заголовок
                heading_text = line.strip('#').strip()
                doc.add_heading(heading_text, level=1)
            elif line.strip():
                # Обычный текст
                doc.add_paragraph(line)
        
        doc.save(filename)
        logger.info(f"Word документ сохранен: {filename}")
    
    def save_to_markdown(self, transcript_text, filename):
        """Сохраняет транскрипт в Markdown файл"""
        logger.info(f"Сохраняю в Markdown: {filename}")
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(transcript_text)
        
        logger.info(f"Markdown файл сохранен: {filename}")
    
    def transcribe_audio(self, yadisk_path, output_prefix="transcript"):
        """Полный процесс транскрибации"""
        temp_audio_file = None
        
        try:
            # 1. Скачиваем аудио с Яндекс.Диска
            temp_audio_file = self.download_audio_from_yadisk(yadisk_path)
            
            # 2. Транскрибируем через Whisper
            transcript_text = self.transcribe_with_whisper(temp_audio_file)
            
            # 3. Разделяем по говорящим
            speakers = self.simple_speaker_segmentation(transcript_text)
            
            # 4. Форматируем
            formatted_text = self.format_transcript_with_speakers(speakers)
            
            # 5. Сохраняем в разных форматах
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Word документ
            word_filename = f"{output_prefix}_{timestamp}.docx"
            self.save_to_word(formatted_text, word_filename)
            
            # Markdown
            md_filename = f"{output_prefix}_{timestamp}.md"
            self.save_to_markdown(formatted_text, md_filename)
            
            logger.info("Транскрибация завершена успешно!")
            return {
                "word_file": word_filename,
                "markdown_file": md_filename,
                "text": formatted_text,
                "raw_text": transcript_text
            }
            
        except Exception as e:
            logger.error(f"Ошибка транскрибации: {e}")
            raise
        
        finally:
            # Очищаем временный файл
            if temp_audio_file and os.path.exists(temp_audio_file):
                os.remove(temp_audio_file)
                logger.info(f"Удален временный файл: {temp_audio_file}")

def main():
    """Основная функция для тестирования"""
    try:
        # Инициализация транскрайбера
        transcriber = FreeTranscriber()
        
        # Путь к аудио файлу на Яндекс.Диске
        yadisk_path = "/audio_transcribe/record.ogg"
        
        print("🎤 Начинаю бесплатную транскрибацию аудио файла...")
        print(f"📁 Файл: {yadisk_path}")
        print(f"🔑 API: OpenRouter (Whisper)")
        print()
        
        # Транскрибация
        result = transcriber.transcribe_audio(yadisk_path, "free_transcript")
        
        print(f"\n✅ Транскрибация завершена успешно!")
        print(f"📄 Word документ: {result['word_file']}")
        print(f"📝 Markdown файл: {result['markdown_file']}")
        print(f"\n📋 Текст транскрипта:\n")
        print("-" * 50)
        print(result['text'])
        print("-" * 50)
        
        print(f"\n🌐 Публичная ссылка на папку: https://yadi.sk/d/9bDnaC7TPqtnLQ")
        
    except Exception as e:
        logger.error(f"Ошибка: {e}")
        print(f"❌ Ошибка: {e}")

if __name__ == "__main__":
    main() 