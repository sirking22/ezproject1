#!/usr/bin/env python3
"""
Транскрибация аудио с разделением по ролям через AssemblyAI
Вывод в Word/Google Docs формат
"""

import os
import requests
import json
import time
from datetime import datetime
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yadisk
import tempfile
import argparse

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AssemblyAITranscriber:
    def __init__(self):
        self.api_key = os.getenv('ASSEMBLYAI_API_KEY')
        if not self.api_key:
            raise ValueError("ASSEMBLYAI_API_KEY не найден в переменных окружения")
        
        self.base_url = "https://api.assemblyai.com/v2"
        self.headers = {
            "authorization": self.api_key,
            "content-type": "application/json"
        }
        
        # Инициализация Яндекс.Диска
        self.yandex_token = os.getenv('YA_ACCESS_TOKEN')
        if self.yandex_token:
            self.yandex = yadisk.YaDisk(token=self.yandex_token)
    
    def download_from_yandex(self, file_path):
        """Скачивает файл с Яндекс.Диска во временную папку"""
        try:
            logger.info(f"Скачиваю файл с Яндекс.Диска: {file_path}")
            
            # Создаем временный файл
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.ogg')
            temp_path = temp_file.name
            temp_file.close()
            
            # Скачиваем файл
            self.yandex.download(file_path, temp_path)
            
            logger.info(f"Файл скачан: {temp_path}")
            return temp_path
            
        except Exception as e:
            logger.error(f"Ошибка скачивания с Яндекс.Диска: {e}")
            raise
    
    def upload_audio(self, audio_path):
        """Загружает аудио файл в AssemblyAI"""
        try:
            logger.info(f"Загружаю аудио: {audio_path}")
            
            # Если это URL, пробуем загрузить напрямую
            if audio_path.startswith('http'):
                upload_url = f"{self.base_url}/upload"
                upload_data = {"url": audio_path}
                response = requests.post(upload_url, json=upload_data, headers=self.headers)
            else:
                # Если это локальный файл, загружаем содержимое
                with open(audio_path, "rb") as f:
                    upload_url = f"{self.base_url}/upload"
                    response = requests.post(upload_url, data=f, headers=self.headers)
            
            response.raise_for_status()
            upload_id = response.json()["upload_url"]
            logger.info(f"Аудио загружено, ID: {upload_id}")
            return upload_id
            
        except Exception as e:
            logger.error(f"Ошибка загрузки аудио: {e}")
            raise
    
    def start_transcription(self, upload_id, expected_speakers=None):
        """Запускает транскрибацию с динамическим распознаванием говорящих"""
        try:
            logger.info("Запускаю транскрибацию с динамическим Speaker Diarization")
            
            transcript_url = f"{self.base_url}/transcript"
            transcript_data = {
                "audio_url": upload_id,
                "speaker_labels": True,
                "language_code": "ru",
                "punctuate": True,
                "format_text": True
            }
            
            # Динамическое определение количества участников
            if expected_speakers:
                transcript_data["speakers_expected"] = expected_speakers
                logger.info(f"Ожидаем {expected_speakers} участников")
            else:
                # Для совещаний с множеством спикеров - не ограничиваем
                logger.info("Динамическое определение количества участников")
            
            response = requests.post(transcript_url, json=transcript_data, headers=self.headers)
            response.raise_for_status()
            
            transcript_id = response.json()["id"]
            logger.info(f"Транскрибация запущена, ID: {transcript_id}")
            return transcript_id
            
        except Exception as e:
            logger.error(f"Ошибка запуска транскрибации: {e}")
            raise
    
    def get_transcription_status(self, transcript_id):
        """Проверяет статус транскрибации"""
        transcript_url = f"{self.base_url}/transcript/{transcript_id}"
        
        response = requests.get(transcript_url, headers=self.headers)
        response.raise_for_status()
        
        return response.json()
    
    def wait_for_completion(self, transcript_id, timeout=300):
        """Ожидает завершения транскрибации"""
        logger.info("Ожидаю завершения транскрибации...")
        
        start_time = time.time()
        while time.time() - start_time < timeout:
            status_data = self.get_transcription_status(transcript_id)
            status = status_data["status"]
            
            if status == "completed":
                logger.info("Транскрибация завершена!")
                return status_data
            elif status == "error":
                error_msg = status_data.get("error", "Неизвестная ошибка")
                raise Exception(f"Ошибка транскрибации: {error_msg}")
            
            logger.info(f"Статус: {status}, ожидаю...")
            time.sleep(10)
        
        raise TimeoutError("Превышено время ожидания транскрибации")
    
    def format_transcript_with_speakers(self, transcript_data):
        """Форматирует транскрипт с разделением по говорящим"""
        logger.info("Форматирую транскрипт по ролям...")
        
        utterances = transcript_data.get("utterances", [])
        if not utterances:
            logger.warning("Нет данных о говорящих, используем обычную транскрибацию")
            return self.format_simple_transcript(transcript_data)
        
        formatted_text = []
        formatted_text.append("# ТРАНСКРИПТ С РАЗДЕЛЕНИЕМ ПО РОЛЯМ\n")
        formatted_text.append(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
        formatted_text.append("=" * 50 + "\n\n")
        
        current_speaker = None
        
        for utterance in utterances:
            speaker = f"Говорящий {utterance['speaker']}"
            text = utterance['text']
            
            if speaker != current_speaker:
                formatted_text.append(f"\n**{speaker}:**\n")
                current_speaker = speaker
            
            formatted_text.append(f"{text}\n")
        
        return "\n".join(formatted_text)
    
    def format_simple_transcript(self, transcript_data):
        """Форматирует обычный транскрипт без разделения по говорящим"""
        text = transcript_data.get("text", "")
        return f"# ТРАНСКРИПТ\n\n{text}"
    
    def save_to_word(self, transcript_text, filename):
        """Сохраняет транскрипт в Word документ"""
        logger.info(f"Сохраняю в Word: {filename}")
        
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ТРАНСКРИПТ С РАЗДЕЛЕНИЕМ ПО РОЛЯМ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата
        date_para = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Разделитель
        doc.add_paragraph("=" * 50)
        doc.add_paragraph()
        
        # Текст транскрипта
        lines = transcript_text.split('\n')
        for line in lines:
            if line.startswith('**') and line.endswith('**'):
                # Говорящий
                speaker = line.strip('*')
                doc.add_heading(speaker, level=2)
            elif line.startswith('#'):
                # Заголовок
                heading_text = line.strip('#').strip()
                doc.add_heading(heading_text, level=1)
            elif line.strip():
                # Обычный текст
                doc.add_paragraph(line)
        
        doc.save(filename)
        logger.info(f"Word документ сохранен: {filename}")
    
    def save_to_markdown(self, transcript_text, filename):
        """Сохраняет транскрипт в Markdown файл (для Google Docs)"""
        logger.info(f"Сохраняю в Markdown: {filename}")
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(transcript_text)
        
        logger.info(f"Markdown файл сохранен: {filename}")
    
    def transcribe_audio(self, audio_path, output_name, expected_speakers=None):
        """Транскрибирует аудио с динамическим распознаванием говорящих"""
        try:
            # 1. Загрузка аудио
            if audio_path.startswith('/'):
                # Путь на Яндекс.Диске
                local_path = self.download_from_yandex(audio_path)
                upload_id = self.upload_audio(local_path)
                # Удаляем временный файл
                os.unlink(local_path)
            else:
                # Локальный файл или URL
                upload_id = self.upload_audio(audio_path)
            
            # 2. Запуск транскрибации с динамическим определением участников
            transcript_id = self.start_transcription(upload_id, expected_speakers)
            
            # 3. Ожидание завершения
            logger.info("Ожидаю завершения транскрибации...")
            while True:
                status = self.get_transcription_status(transcript_id)
                if status == "completed":
                    break
                elif status == "error":
                    raise Exception("Ошибка транскрибации")
                time.sleep(5)
            
            # 4. Получение результата
            transcript_data = self.get_transcription_result(transcript_id)
            
            # 5. Создание документов
            result = self.create_documents(transcript_data, output_name)
            
            logger.info(f"✅ Транскрибация завершена! Найдено {result['speakers_count']} говорящих")
            return result
            
        except Exception as e:
            logger.error(f"Ошибка транскрибации: {e}")
            raise
    
    def analyze_speakers(self, transcript_data):
        """Анализирует и улучшает распознавание говорящих"""
        try:
            speakers = {}
            
            # Группируем слова по говорящим
            for word in transcript_data.get("words", []):
                speaker = word.get("speaker", "Unknown")
                if speaker not in speakers:
                    speakers[speaker] = {
                        "words": [],
                        "total_duration": 0,
                        "segments": []
                    }
                
                speakers[speaker]["words"].append(word)
                duration = word.get("end", 0) - word.get("start", 0)
                speakers[speaker]["total_duration"] += duration
            
            # Создаем сегменты для каждого говорящего
            for speaker, data in speakers.items():
                current_segment = {"text": "", "start": None, "end": None}
                
                for word in data["words"]:
                    if current_segment["start"] is None:
                        current_segment["start"] = word.get("start", 0)
                    
                    current_segment["text"] += word.get("text", "") + " "
                    current_segment["end"] = word.get("end", 0)
                    
                    # Если пауза больше 2 секунд, создаем новый сегмент
                    if len(data["words"]) > 1:
                        next_word_idx = data["words"].index(word) + 1
                        if next_word_idx < len(data["words"]):
                            next_word = data["words"][next_word_idx]
                            if next_word.get("start", 0) - word.get("end", 0) > 2000:
                                if current_segment["text"].strip():
                                    data["segments"].append(current_segment)
                                current_segment = {"text": "", "start": None, "end": None}
                
                # Добавляем последний сегмент
                if current_segment["text"].strip():
                    data["segments"].append(current_segment)
            
            # Сортируем говорящих по длительности речи
            sorted_speakers = sorted(speakers.items(), 
                                   key=lambda x: x[1]["total_duration"], 
                                   reverse=True)
            
            # Переназначаем говорящих (Speaker A, B, C, D, E, F, G, H, I, J...)
            speaker_mapping = {}
            for i, (old_speaker, data) in enumerate(sorted_speakers):
                if i < 26:  # A-Z (26 букв)
                    new_speaker = f"Speaker {chr(65 + i)}"  # A, B, C, D, E, F, G, H, I, J...
                else:
                    new_speaker = f"Speaker {i+1}"  # Speaker 27, Speaker 28...
                speaker_mapping[old_speaker] = new_speaker
            
            logger.info(f"Найдено {len(speakers)} говорящих:")
            for i, (speaker, data) in enumerate(sorted_speakers):
                new_name = speaker_mapping[speaker]
                logger.info(f"  {new_name}: {data['total_duration']/1000:.1f}с, "
                           f"{len(data['words'])} слов, "
                           f"{len(data['segments'])} сегментов")
            
            return speaker_mapping, speakers
            
        except Exception as e:
            logger.error(f"Ошибка анализа говорящих: {e}")
            return {}, {}
    
    def create_documents(self, transcript_data, output_name):
        """Создает Word и Markdown документы с корректным порядком реплик (по utterances) в папке transcripts/"""
        try:
            # Создаём папку transcripts, если нет
            transcripts_dir = "transcripts"
            os.makedirs(transcripts_dir, exist_ok=True)

            # Анализируем говорящих для статистики
            speaker_mapping, speakers_analysis = self.analyze_speakers(transcript_data)

            # Создаем Word документ
            doc = Document()
            doc.add_heading(f"Транскрипт: {output_name}", 0)

            # Добавляем статистику говорящих
            doc.add_heading("Статистика говорящих", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Говорящий"
            hdr_cells[1].text = "Время (сек)"
            hdr_cells[2].text = "Слов"
            hdr_cells[3].text = "Сегментов"

            for speaker, data in speakers_analysis.items():
                new_name = speaker_mapping.get(speaker, speaker)
                row_cells = table.add_row().cells
                row_cells[0].text = new_name
                row_cells[1].text = f"{data['total_duration']/1000:.1f}"
                row_cells[2].text = str(len(data['words']))
                row_cells[3].text = f"{len(data['segments'])} сегментов"

            doc.add_heading("Транскрипт", level=1)

            # Используем только utterances для реального порядка
            utterances = transcript_data.get("utterances", [])
            full_text = ""
            for utt in utterances:
                speaker = f"Speaker {utt['speaker']}"
                text = utt['text'].strip()
                # Word
                p = doc.add_paragraph()
                p.add_run(f"**{speaker}:** ").bold = True
                p.add_run(text)
                # Markdown
                full_text += f"**{speaker}:** {text}\n\n"

            # Сохраняем Word документ
            word_file = os.path.join(transcripts_dir, f"{output_name}.docx")
            doc.save(word_file)

            # Сохраняем Markdown
            markdown_file = os.path.join(transcripts_dir, f"{output_name}.md")
            with open(markdown_file, "w", encoding="utf-8") as f:
                f.write(f"# Транскрипт: {output_name}\n\n")
                f.write("## Статистика говорящих\n\n")
                f.write("| Говорящий | Время (сек) | Слов | Сегментов |\n")
                f.write("|-----------|-------------|------|-----------|\n")
                for speaker, data in speakers_analysis.items():
                    new_name = speaker_mapping.get(speaker, speaker)
                    f.write(f"| {new_name} | {data['total_duration']/1000:.1f} | {len(data['words'])} | {len(data['segments'])} |\n")
                f.write("\n## Диалог\n\n")
                f.write(full_text)

            return {
                "word_file": word_file,
                "markdown_file": markdown_file,
                "text": full_text,
                "speakers_count": len(speakers_analysis)
            }
        except Exception as e:
            logger.error(f"Ошибка создания документов: {e}")
            raise

    def get_transcription_result(self, transcript_id):
        """Получает результат транскрибации"""
        try:
            transcript_url = f"{self.base_url}/transcript/{transcript_id}"
            response = requests.get(transcript_url, headers=self.headers)
            response.raise_for_status()
            
            return response.json()
            
        except Exception as e:
            logger.error(f"Ошибка получения результата: {e}")
            raise

def main():
    parser = argparse.ArgumentParser(description="Транскрибация аудио с динамическим определением говорящих (AssemblyAI)")
    parser.add_argument("audio_path", nargs="?", help="Путь к аудиофайлу (локально или на Яндекс.Диске)")
    parser.add_argument("output_name", nargs="?", help="Имя для результата (без расширения)")
    parser.add_argument("--expected_speakers", type=int, default=None, help="Ожидаемое число говорящих (опционально)")
    parser.add_argument("--fetch-only", type=str, default=None, help="ID транскрипции для скачивания результата без ожидания")
    args = parser.parse_args()

    transcriber = AssemblyAITranscriber()

    if args.fetch_only:
        transcript_id = args.fetch_only
        print(f"🔎 Получаю результат транскрипции по ID: {transcript_id}")
        transcript_data = transcriber.get_transcription_result(transcript_id)
        result = transcriber.create_documents(transcript_data, args.output_name or f"transcript_{transcript_id}")
        print(f"\n✅ Результат получен!")
        print(f"👥 Найдено говорящих: {result.get('speakers_count', '?')}")
        print(f"📄 Word: {result['word_file']}")
        print(f"📝 Markdown: {result['markdown_file']}\n")
        print(f"\n📋 Текст транскрипта:\n{result['text'][:1000]}{'...' if len(result['text'])>1000 else ''}")
        return

    if not args.audio_path or not args.output_name:
        parser.error("audio_path и output_name обязательны, если не используется --fetch-only")

    result = transcriber.transcribe_audio(args.audio_path, args.output_name, expected_speakers=args.expected_speakers)
    print(f"\n✅ Транскрибация завершена!")
    print(f"👥 Найдено говорящих: {result.get('speakers_count', '?')}")
    print(f"📄 Word: {result['word_file']}")
    print(f"📝 Markdown: {result['markdown_file']}")
    print(f"\n📋 Текст транскрипта:\n{result['text'][:1000]}{'...' if len(result['text'])>1000 else ''}")

if __name__ == "__main__":
    main() 