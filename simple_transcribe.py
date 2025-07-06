#!/usr/bin/env python3
"""
Упрощенная транскрибация с базовым распознаванием говорящих
"""

import os
import requests
import json
import time
import logging
from docx import Document
from dotenv import load_dotenv
import yadisk
import tempfile

# Загружаем .env файл
load_dotenv()

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class SimpleTranscriber:
    def __init__(self):
        self.api_key = os.getenv('ASSEMBLYAI_API_KEY')
        if not self.api_key:
            raise ValueError("ASSEMBLYAI_API_KEY не найден в переменных окружения")
        
        self.base_url = "https://api.assemblyai.com/v2"
        self.headers = {
            "authorization": self.api_key,
            "content-type": "application/json"
        }
        
        # Инициализация Яндекс.Диска
        self.yandex_token = os.getenv('YA_ACCESS_TOKEN')
        if self.yandex_token:
            self.yandex = yadisk.YaDisk(token=self.yandex_token)
    
    def download_from_yandex(self, file_path):
        """Скачивает файл с Яндекс.Диска"""
        try:
            logger.info(f"Скачиваю файл с Яндекс.Диска: {file_path}")
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.ogg')
            temp_path = temp_file.name
            temp_file.close()
            
            self.yandex.download(file_path, temp_path)
            logger.info(f"Файл скачан: {temp_path}")
            return temp_path
            
        except Exception as e:
            logger.error(f"Ошибка скачивания: {e}")
            raise
    
    def upload_audio(self, audio_path):
        """Загружает аудио в AssemblyAI"""
        try:
            logger.info(f"Загружаю аудио: {audio_path}")
            
            with open(audio_path, "rb") as f:
                upload_url = f"{self.base_url}/upload"
                response = requests.post(upload_url, data=f, headers=self.headers)
            
            response.raise_for_status()
            upload_id = response.json()["upload_url"]
            logger.info(f"Аудио загружено, ID: {upload_id}")
            return upload_id
            
        except Exception as e:
            logger.error(f"Ошибка загрузки: {e}")
            raise
    
    def start_transcription(self, upload_id):
        """Запускает транскрибацию"""
        try:
            logger.info("Запускаю транскрибацию...")
            
            transcript_url = f"{self.base_url}/transcript"
            transcript_data = {
                "audio_url": upload_id,
                "speaker_labels": True,
                "language_code": "ru",
                "punctuate": True,
                "format_text": True
            }
            
            response = requests.post(transcript_url, json=transcript_data, headers=self.headers)
            response.raise_for_status()
            
            transcript_id = response.json()["id"]
            logger.info(f"Транскрибация запущена, ID: {transcript_id}")
            return transcript_id
            
        except Exception as e:
            logger.error(f"Ошибка запуска: {e}")
            raise
    
    def get_transcription_status(self, transcript_id):
        """Получает статус транскрибации"""
        try:
            transcript_url = f"{self.base_url}/transcript/{transcript_id}"
            response = requests.get(transcript_url, headers=self.headers)
            response.raise_for_status()
            
            return response.json()["status"]
            
        except Exception as e:
            logger.error(f"Ошибка получения статуса: {e}")
            raise
    
    def get_transcription_result(self, transcript_id):
        """Получает результат транскрибации"""
        try:
            transcript_url = f"{self.base_url}/transcript/{transcript_id}"
            response = requests.get(transcript_url, headers=self.headers)
            response.raise_for_status()
            
            return response.json()
            
        except Exception as e:
            logger.error(f"Ошибка получения результата: {e}")
            raise
    
    def create_documents(self, transcript_data, output_name):
        """Создает документы"""
        try:
            # Анализируем говорящих
            speakers = {}
            for word in transcript_data.get("words", []):
                speaker = word.get("speaker", "Unknown")
                if speaker not in speakers:
                    speakers[speaker] = {"words": [], "total_duration": 0}
                
                speakers[speaker]["words"].append(word)
                duration = word.get("end", 0) - word.get("start", 0)
                speakers[speaker]["total_duration"] += duration
            
            # Сортируем говорящих по времени
            sorted_speakers = sorted(speakers.items(), 
                                   key=lambda x: x[1]["total_duration"], 
                                   reverse=True)
            
            # Создаем маппинг
            speaker_mapping = {}
            for i, (old_speaker, data) in enumerate(sorted_speakers):
                new_speaker = f"Speaker {chr(65 + i)}"
                speaker_mapping[old_speaker] = new_speaker
            
            # Создаем Word документ
            doc = Document()
            doc.add_heading(f"Транскрипт: {output_name}", 0)
            
            # Статистика
            doc.add_heading("Статистика говорящих", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Говорящий"
            hdr_cells[1].text = "Время (сек)"
            hdr_cells[2].text = "Слов"
            
            for speaker, data in sorted_speakers:
                new_name = speaker_mapping[speaker]
                row_cells = table.add_row().cells
                row_cells[0].text = new_name
                row_cells[1].text = f"{data['total_duration']/1000:.1f}"
                row_cells[2].text = str(len(data['words']))
            
            doc.add_heading("Транскрипт", level=1)
            
            # Транскрипт
            full_text = ""
            current_speaker = None
            current_text = ""
            
            for word in transcript_data.get("words", []):
                speaker = word.get("speaker", "Unknown")
                text = word.get("text", "")
                
                if speaker != current_speaker:
                    if current_text:
                        new_speaker = speaker_mapping.get(current_speaker, current_speaker)
                        p = doc.add_paragraph()
                        p.add_run(f"**{new_speaker}:** ").bold = True
                        p.add_run(current_text.strip())
                        full_text += f"**{new_speaker}:** {current_text.strip()}\n\n"
                    
                    current_speaker = speaker
                    current_text = text
                else:
                    current_text += " " + text
            
            # Последний сегмент
            if current_text:
                new_speaker = speaker_mapping.get(current_speaker, current_speaker)
                p = doc.add_paragraph()
                p.add_run(f"**{new_speaker}:** ").bold = True
                p.add_run(current_text.strip())
                full_text += f"**{new_speaker}:** {current_text.strip()}\n\n"
            
            # Сохраняем файлы
            word_file = f"{output_name}.docx"
            doc.save(word_file)
            
            markdown_file = f"{output_name}.md"
            with open(markdown_file, "w", encoding="utf-8") as f:
                f.write(f"# Транскрипт: {output_name}\n\n")
                f.write("## Статистика говорящих\n\n")
                f.write("| Говорящий | Время (сек) | Слов |\n")
                f.write("|-----------|-------------|------|\n")
                
                for speaker, data in sorted_speakers:
                    new_name = speaker_mapping[speaker]
                    f.write(f"| {new_name} | {data['total_duration']/1000:.1f} | {len(data['words'])} |\n")
                
                f.write(f"\n## Транскрипт\n\n{full_text}")
            
            logger.info(f"Документы созданы: {word_file}, {markdown_file}")
            return {
                "word_file": word_file,
                "markdown_file": markdown_file,
                "text": full_text,
                "speakers_count": len(speakers)
            }
            
        except Exception as e:
            logger.error(f"Ошибка создания документов: {e}")
            raise
    
    def transcribe_audio(self, audio_path, output_name):
        """Полный процесс транскрибации"""
        try:
            # 1. Загрузка аудио
            if audio_path.startswith('/'):
                local_path = self.download_from_yandex(audio_path)
                upload_id = self.upload_audio(local_path)
                os.unlink(local_path)
            else:
                upload_id = self.upload_audio(audio_path)
            
            # 2. Запуск транскрибации
            transcript_id = self.start_transcription(upload_id)
            
            # 3. Ожидание завершения
            logger.info("Ожидаю завершения транскрибации...")
            while True:
                status = self.get_transcription_status(transcript_id)
                if status == "completed":
                    break
                elif status == "error":
                    raise Exception("Ошибка транскрибации")
                time.sleep(5)
            
            # 4. Получение результата
            transcript_data = self.get_transcription_result(transcript_id)
            
            # 5. Создание документов
            result = self.create_documents(transcript_data, output_name)
            
            logger.info(f"✅ Транскрибация завершена! Найдено {result['speakers_count']} говорящих")
            return result
            
        except Exception as e:
            logger.error(f"Ошибка транскрибации: {e}")
            raise

def main():
    """Тестирование"""
    try:
        transcriber = SimpleTranscriber()
        result = transcriber.transcribe_audio("/audio_transcribe/record.ogg", "simple_transcript")
        
        print(f"\n✅ Транскрибация завершена!")
        print(f"👥 Найдено говорящих: {result['speakers_count']}")
        print(f"📄 Word: {result['word_file']}")
        print(f"📝 Markdown: {result['markdown_file']}")
        
    except Exception as e:
        print(f"❌ Ошибка: {e}")

if __name__ == "__main__":
    main() 